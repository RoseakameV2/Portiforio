# NÃO ESQUECE KEEPLINE, VOCÊ É O MELHOR PROGRAMADOR DO MUNDO!

#  C:\Python314\python.exe -m PyInstaller --onefile --windowed "C:\Python\Boletos AUTO\Módulos internos\app.py" 

import customtkinter as ctk
from tkinter import messagebox
from docx import Document
from docx.shared import RGBColor
import win32api
import win32print
import win32com.client
import os
import sys
import time
import shutil


ctk.set_appearance_mode("light")
ctk.set_default_color_theme("blue")

if getattr(sys, 'frozen', False):
    BASE_DIR = os.path.dirname(sys.argv[0])
else:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))

PASTA_MODELOS = os.path.join(BASE_DIR, "Models")
PASTA_SAIDA = os.path.join(BASE_DIR, "Gerados")
PASTA_TEMP = os.path.join(BASE_DIR, "_temp_doc")

os.makedirs(PASTA_MODELOS, exist_ok=True)
os.makedirs(PASTA_SAIDA, exist_ok=True)
os.makedirs(PASTA_TEMP, exist_ok=True)

MESES = [
    "JANEIRO", "FEVEREIRO", "MARÇO", "ABRIL", "MAIO", "JUNHO",
    "JULHO", "AGOSTO", "SETEMBRO", "OUTUBRO", "NOVEMBRO", "DEZEMBRO"
]


def buscar_modelos():
    return sorted([
        f for f in os.listdir(PASTA_MODELOS)
        if (f.lower().endswith(".docx") or f.lower().endswith(".doc")) and not f.startswith("~$")
    ])

def listar_impressoras():
    try:
        return [
            p[2] for p in win32print.EnumPrinters(
                win32print.PRINTER_ENUM_LOCAL | win32print.PRINTER_ENUM_CONNECTIONS
            )
        ]
    except:
        return ["Padrão"]

def converter_doc_para_docx(caminho_doc):
    nome = os.path.splitext(os.path.basename(caminho_doc))[0]
    destino = os.path.join(PASTA_TEMP, f"{nome}.docx")

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc = word.Documents.Open(caminho_doc)
    doc.SaveAs(destino, FileFormat=16)
    doc.Close()
    word.Quit()

    return destino

def substituir_texto(paragrafos, substituicoes):
    for p in paragrafos:
        texto_original = p.text
        texto_novo = texto_original
        for k, v in substituicoes.items():
            texto_novo = texto_novo.replace(k, v)

        if texto_novo != texto_original:
            p.clear()
            run = p.add_run(texto_novo)
            run.font.color.rgb = RGBColor(0, 0, 0)


def executar():
    mes_nome = combo_mes.get()
    ano = entry_ano.get()
    imprimir = var_imprimir.get()
    marcar_todos = var_todos.get()

    if not mes_nome or not ano:
        messagebox.showerror("Erro", "Informe o mês e o ano.")
        return

    arquivos = list(dict_checks.keys()) if marcar_todos else [
        a for a, v in dict_checks.items() if v.get() == 1
    ]

    if not arquivos:
        messagebox.showerror("Erro", "Selecione ao menos um modelo.")
        return

    total = len(arquivos)
    progresso.set(0)
    app.update_idletasks()

    mes_num = str(MESES.index(mes_nome) + 1).zfill(2)
    substituicoes = {
        "{{MES_NOME}}": mes_nome,
        "{{MES_NUM}}": mes_num,
        "{{ano}}": ano
    }

    if imprimir:
        try:
            win32print.SetDefaultPrinter(combo_impressoras.get())
        except:
            pass

    erros = []
    sucessos = 0

    for i, nome in enumerate(arquivos, start=1):
        try:
            caminho_original = os.path.join(PASTA_MODELOS, nome)

            if nome.lower().endswith(".doc"):
                caminho_trabalho = converter_doc_para_docx(caminho_original)
            else:
                caminho_trabalho = caminho_original

            doc = Document(caminho_trabalho)

            substituir_texto(doc.paragraphs, substituicoes)
            for tabela in doc.tables:
                for linha in tabela.rows:
                    for celula in linha.cells:
                        substituir_texto(celula.paragraphs, substituicoes)

            nome_base = os.path.splitext(nome)[0]
            saida = os.path.join(PASTA_SAIDA, f"{nome_base}_{mes_nome}.docx")
            doc.save(saida)

            if imprimir:
                win32api.ShellExecute(0, "print", saida, None, ".", 0)
                time.sleep(3)  # impressão 1 a 1 (anti-travamento)

            sucessos += 1

        except Exception as e:
            erros.append(f"{nome}: {e}")

        
        progresso.set(i / total)
        app.update_idletasks()

    shutil.rmtree(PASTA_TEMP, ignore_errors=True)

    msg = f"Finalizado!\nSucessos: {sucessos}"
    if erros:
        msg += "\n\nErros:\n" + "\n".join(erros)
        messagebox.showwarning("Relatório", msg)
    else:
        messagebox.showinfo("Sucesso", msg)


app = ctk.CTk()
app.title("Gerador de Recibos - By keepline")
app.geometry("600x720")
app.minsize(500, 520)

ctk.CTkLabel(app, text="GERADOR DE RECIBOS", font=("Roboto", 22, "bold")).pack(pady=15)

ctk.CTkLabel(app, text="Selecione os modelos:", font=("Roboto", 12, "bold")).pack(anchor="w", padx=30)

frame_lista = ctk.CTkScrollableFrame(app)
frame_lista.pack(padx=30, pady=10, fill="both", expand=True)

dict_checks = {}
for arq in buscar_modelos():
    var = ctk.IntVar(value=0)
    chk = ctk.CTkCheckBox(frame_lista, text=arq, variable=var)
    chk.pack(anchor="w", padx=10, pady=4, fill="x")
    dict_checks[arq] = var

frame_cfg = ctk.CTkFrame(app)
frame_cfg.pack(padx=30, pady=10, fill="x")

var_todos = ctk.CTkCheckBox(frame_cfg, text="Marcar todos")
var_todos.pack(anchor="w", padx=10, pady=5)

ctk.CTkLabel(frame_cfg, text="Mês:").pack(anchor="w", padx=10)
combo_mes = ctk.CTkOptionMenu(frame_cfg, values=MESES)
combo_mes.pack(fill="x", padx=10, pady=5)

ctk.CTkLabel(frame_cfg, text="Ano:").pack(anchor="w", padx=10)
entry_ano = ctk.CTkEntry(frame_cfg)
entry_ano.insert(0, "2026")
entry_ano.pack(fill="x", padx=10, pady=5)

ctk.CTkLabel(frame_cfg, text="Impressora:").pack(anchor="w", padx=10)
combo_impressoras = ctk.CTkOptionMenu(frame_cfg, values=listar_impressoras())
combo_impressoras.pack(fill="x", padx=10, pady=5)

var_imprimir = ctk.CTkCheckBox(frame_cfg, text="Imprimir automaticamente")
var_imprimir.pack(anchor="w", padx=10, pady=5)


progresso = ctk.CTkProgressBar(app)
progresso.set(0)
progresso.pack(padx=30, pady=(10, 5), fill="x")


ctk.CTkButton(
    app,
    text="GERAR RECIBOS",
    command=executar,
    height=50,
    font=("Roboto", 16, "bold"),
    fg_color="#27ae60",
    hover_color="#219150"
).pack(padx=30, pady=20, fill="x")

app.mainloop()
